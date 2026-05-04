from __future__ import annotations

import argparse
import json
import math
import os
import re
import sys
import tempfile
import urllib.request
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from faster_whisper import WhisperModel


USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/135.0 Safari/537.36"
)


@dataclass
class EpisodeManifest:
    title: str
    published_at: str
    published_label: str
    duration: str
    file_name: str
    episode_url: str
    audio_url: str


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Transcribe podcast episode audio from a manifest and export transcript docs."
    )
    parser.add_argument("manifest_path", help="Path to manifest.json from the podcast word export step.")
    parser.add_argument(
        "output_dir",
        nargs="?",
        help="Directory for transcript outputs. Defaults to a sibling 'transcripts' folder next to the manifest.",
    )
    parser.add_argument(
        "--model",
        default="base.en",
        help="Whisper model name. Use tiny.en for speed, base.en for a quick quality/speed balance.",
    )
    parser.add_argument(
        "--device",
        default="auto",
        help="Execution device for faster-whisper, usually 'auto', 'cpu', or 'cuda'.",
    )
    parser.add_argument(
        "--compute-type",
        default="int8",
        help="Compute type for faster-whisper, such as int8 or float16.",
    )
    parser.add_argument(
        "--limit",
        type=int,
        default=0,
        help="Optional episode limit for testing. 0 means all episodes.",
    )
    parser.add_argument(
        "--overwrite",
        action="store_true",
        help="Overwrite existing transcript outputs instead of skipping them.",
    )
    parser.add_argument(
        "--beam-size",
        type=int,
        default=5,
        help="Beam size for transcription decoding.",
    )
    parser.add_argument(
        "--vad",
        action="store_true",
        default=True,
        help="Enable voice activity detection filtering.",
    )
    parser.add_argument(
        "--no-vad",
        dest="vad",
        action="store_false",
        help="Disable voice activity detection filtering.",
    )
    return parser.parse_args()


def normalize_whitespace(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\xa0", " ")).strip()


def sanitize_file_stem(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "", value)
    cleaned = normalize_whitespace(cleaned).rstrip(". ")
    return cleaned or "untitled"


def format_date(date_value: str) -> str:
    if not date_value:
      return ""

    try:
        parsed = datetime.fromisoformat(date_value)
    except ValueError:
        return date_value

    return parsed.strftime("%B %-d, %Y") if os.name != "nt" else parsed.strftime("%B %#d, %Y")


def load_manifest(path: Path) -> tuple[str, list[EpisodeManifest]]:
    payload = json.loads(path.read_text(encoding="utf-8"))
    show_title = payload.get("showTitle", "Podcast")
    episodes = []

    for item in payload.get("episodes", []):
        episodes.append(
            EpisodeManifest(
                title=item.get("title", "").strip(),
                published_at=item.get("publishedAt", "").strip(),
                published_label=item.get("publishedLabel", "").strip(),
                duration=item.get("duration", "").strip(),
                file_name=item.get("fileName", "").strip(),
                episode_url=item.get("episodeUrl", "").strip(),
                audio_url=item.get("audioUrl", "").strip(),
            )
        )

    return show_title, episodes


def download_audio(audio_url: str, destination: Path) -> None:
    request = urllib.request.Request(
        audio_url,
        headers={
            "User-Agent": USER_AGENT,
            "Accept": "*/*",
        },
    )
    with urllib.request.urlopen(request) as response, destination.open("wb") as output:
        while True:
            chunk = response.read(1024 * 1024)
            if not chunk:
                break
            output.write(chunk)


def seconds_to_timestamp(value: float) -> str:
    total_seconds = max(0, int(math.floor(value)))
    hours = total_seconds // 3600
    minutes = (total_seconds % 3600) // 60
    seconds = total_seconds % 60
    return f"{hours:02}:{minutes:02}:{seconds:02}"


def group_segments(segments: Iterable[dict]) -> list[str]:
    paragraphs: list[str] = []
    current_lines: list[str] = []
    current_length = 0
    previous_end = 0.0

    for segment in segments:
        text = normalize_whitespace(segment["text"])
        if not text:
            continue

        timestamp = seconds_to_timestamp(segment["start"])
        line = f"[{timestamp}] {text}"
        gap = segment["start"] - previous_end

        if current_lines and (gap > 2.0 or current_length > 850):
            paragraphs.append(" ".join(current_lines))
            current_lines = []
            current_length = 0

        current_lines.append(line)
        current_length += len(line)
        previous_end = segment["end"]

        if re.search(r"[.!?][\"']?$", text) and current_length > 260:
            paragraphs.append(" ".join(current_lines))
            current_lines = []
            current_length = 0

    if current_lines:
        paragraphs.append(" ".join(current_lines))

    return paragraphs


def write_outputs(
    output_dir: Path,
    show_title: str,
    episode: EpisodeManifest,
    transcript_text: str,
    segments: list[dict],
    model_name: str,
) -> None:
    base_stem = Path(episode.file_name).stem if episode.file_name else sanitize_file_stem(episode.title)
    txt_path = output_dir / f"{base_stem}.txt"
    json_path = output_dir / f"{base_stem}.segments.json"
    docx_path = output_dir / f"{base_stem}.docx"

    txt_path.write_text(transcript_text, encoding="utf-8")
    json_path.write_text(
        json.dumps(
            {
                "showTitle": show_title,
                "title": episode.title,
                "publishedAt": episode.published_at,
                "duration": episode.duration,
                "episodeUrl": episode.episode_url,
                "audioUrl": episode.audio_url,
                "model": model_name,
                "segments": segments,
            },
            indent=2,
        ),
        encoding="utf-8",
    )

    paragraphs = group_segments(segments)

    document = Document()
    document.add_heading(episode.title, level=1)

    metadata = [
        ("Podcast", show_title),
        ("Published", episode.published_label or format_date(episode.published_at)),
        ("Duration", episode.duration),
        ("Episode URL", episode.episode_url),
        ("Audio URL", episode.audio_url),
        ("Transcription model", model_name),
    ]

    for label, value in metadata:
        if not value:
            continue
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(value)

    document.add_paragraph()
    document.add_heading("Transcript", level=2)

    for paragraph_text in paragraphs:
        document.add_paragraph(paragraph_text)

    document.save(docx_path)


def transcribe_episode(
    model: WhisperModel,
    show_title: str,
    episode: EpisodeManifest,
    output_dir: Path,
    overwrite: bool,
    beam_size: int,
    vad: bool,
    model_name: str,
) -> str:
    base_stem = Path(episode.file_name).stem if episode.file_name else sanitize_file_stem(episode.title)
    docx_path = output_dir / f"{base_stem}.docx"

    if docx_path.exists() and not overwrite:
        return f"Skipped existing transcript: {docx_path.name}"

    if not episode.audio_url:
        return f"Skipped {episode.title}: no audio URL in manifest."

    with tempfile.TemporaryDirectory(prefix="podcast-audio-") as temp_dir:
        audio_path = Path(temp_dir) / "episode.mp3"
        download_audio(episode.audio_url, audio_path)

        segments_iterable, info = model.transcribe(
            str(audio_path),
            beam_size=beam_size,
            language="en",
            vad_filter=vad,
            condition_on_previous_text=False,
        )

        segments = [
            {
                "id": index,
                "start": round(segment.start, 3),
                "end": round(segment.end, 3),
                "text": normalize_whitespace(segment.text),
            }
            for index, segment in enumerate(segments_iterable, start=1)
        ]

        transcript_text = "\n\n".join(group_segments(segments))
        model_used = f"{model_name} ({info.language}, p={info.language_probability:.3f})"
        write_outputs(output_dir, show_title, episode, transcript_text, segments, model_used)

    return f"Transcribed: {episode.title}"


def main() -> int:
    args = parse_args()
    manifest_path = Path(args.manifest_path).resolve()

    if not manifest_path.exists():
        print(f"Manifest not found: {manifest_path}", file=sys.stderr)
        return 1

    show_title, episodes = load_manifest(manifest_path)
    if not episodes:
        print("No episodes found in the manifest.", file=sys.stderr)
        return 1

    if args.limit > 0:
        episodes = episodes[: args.limit]

    output_dir = (
        Path(args.output_dir).resolve()
        if args.output_dir
        else manifest_path.parent / "transcripts"
    )
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"Loading Whisper model '{args.model}' on device '{args.device}'...")
    model = WhisperModel(args.model, device=args.device, compute_type=args.compute_type)

    for episode in episodes:
        message = transcribe_episode(
            model=model,
            show_title=show_title,
            episode=episode,
            output_dir=output_dir,
            overwrite=args.overwrite,
            beam_size=args.beam_size,
            vad=args.vad,
            model_name=args.model,
        )
        print(message)

    print(f"Transcript outputs are in: {output_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
